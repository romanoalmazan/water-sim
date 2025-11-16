#!/usr/bin/env python3
"""
Convert markdown document to Word and PDF formats
Requires: pip install python-docx markdown beautifulsoup4 weasyprint
"""

import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from bs4 import BeautifulSoup
import re
from pathlib import Path
import sys

def markdown_to_word(md_file_path, output_path):
    """Convert markdown file to Word document"""
    
    # Read markdown content
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Convert markdown to HTML
    html = markdown.markdown(md_content, extensions=['codehilite', 'tables', 'toc'])
    soup = BeautifulSoup(html, 'html.parser')
    
    # Create Word document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add title
    title = doc.add_heading('🌱 Predictive Smart Irrigation System - In-Depth Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add table of contents
    doc.add_heading('Table of Contents', level=1)
    toc_p = doc.add_paragraph()
    toc_content = [
        "1. Required Knowledge Domains",
        "2. Potential Questions & Detailed Answers", 
        "3. Learning Path for Building Similar Systems",
        "4. Real-World Implementation Considerations"
    ]
    for item in toc_content:
        toc_p.add_run(item + '\n')
    
    doc.add_page_break()
    
    # Process HTML elements
    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'pre', 'code', 'ul', 'ol']):
        if element.name in ['h1', 'h2', 'h3', 'h4']:
            level = int(element.name[1])
            text = element.get_text().strip()
            # Remove emoji from headings for cleaner Word formatting
            text = re.sub(r'[^\w\s-]', '', text).strip()
            doc.add_heading(text, level=level)
            
        elif element.name == 'p':
            text = element.get_text().strip()
            if text:  # Only add non-empty paragraphs
                p = doc.add_paragraph(text)
                
        elif element.name == 'pre':
            # Code blocks
            code_text = element.get_text()
            p = doc.add_paragraph(code_text)
            p.style = 'Normal'  # You can create a custom code style
            
        elif element.name in ['ul', 'ol']:
            # Lists
            for li in element.find_all('li'):
                text = li.get_text().strip()
                if text:
                    p = doc.add_paragraph(text, style='List Bullet' if element.name == 'ul' else 'List Number')
    
    # Save the document
    doc.save(output_path)
    print(f"Word document saved to: {output_path}")

def markdown_to_pdf_html(md_file_path, html_output_path, pdf_output_path):
    """Convert markdown to PDF via HTML (alternative method)"""
    
    # Read markdown content
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Convert markdown to HTML with better styling
    html = markdown.markdown(md_content, extensions=['codehilite', 'tables', 'toc'])
    
    # Create styled HTML
    styled_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <title>Predictive Smart Irrigation System - Analysis</title>
        <style>
            body {{
                font-family: 'Arial', sans-serif;
                line-height: 1.6;
                max-width: 800px;
                margin: 0 auto;
                padding: 20px;
                color: #333;
            }}
            h1, h2, h3, h4 {{
                color: #2c5aa0;
                margin-top: 30px;
                margin-bottom: 15px;
            }}
            h1 {{
                border-bottom: 3px solid #2c5aa0;
                padding-bottom: 10px;
                text-align: center;
            }}
            h2 {{
                border-bottom: 1px solid #2c5aa0;
                padding-bottom: 5px;
            }}
            code {{
                background-color: #f4f4f4;
                padding: 2px 5px;
                border-radius: 3px;
                font-family: 'Courier New', monospace;
            }}
            pre {{
                background-color: #f4f4f4;
                padding: 15px;
                border-radius: 5px;
                overflow-x: auto;
                border-left: 4px solid #2c5aa0;
            }}
            blockquote {{
                border-left: 4px solid #2c5aa0;
                margin-left: 0;
                padding-left: 20px;
                color: #666;
                font-style: italic;
            }}
            table {{
                border-collapse: collapse;
                width: 100%;
                margin: 20px 0;
            }}
            th, td {{
                border: 1px solid #ddd;
                padding: 8px;
                text-align: left;
            }}
            th {{
                background-color: #f2f2f2;
                font-weight: bold;
            }}
            .toc {{
                background-color: #f9f9f9;
                padding: 20px;
                border-radius: 5px;
                margin: 20px 0;
            }}
            @page {{
                margin: 1in;
            }}
        </style>
    </head>
    <body>
        {html}
    </body>
    </html>
    """
    
    # Save HTML file
    with open(html_output_path, 'w', encoding='utf-8') as f:
        f.write(styled_html)
    
    print(f"HTML file saved to: {html_output_path}")
    
    # Try to convert to PDF using weasyprint
    try:
        import weasyprint
        weasyprint.HTML(string=styled_html).write_pdf(pdf_output_path)
        print(f"PDF file saved to: {pdf_output_path}")
    except ImportError:
        print("WeasyPrint not installed. Install with: pip install weasyprint")
        print("Alternatively, you can convert the HTML file to PDF using your browser:")
        print(f"Open {html_output_path} in browser and use 'Print to PDF'")

def main():
    # Set up paths
    base_dir = Path(__file__).parent
    docs_dir = base_dir / "docs"
    docs_dir.mkdir(exist_ok=True)
    
    md_file = docs_dir / "Predictive_Smart_Irrigation_Analysis.md"
    word_file = docs_dir / "Predictive_Smart_Irrigation_Analysis.docx"
    html_file = docs_dir / "Predictive_Smart_Irrigation_Analysis.html"
    pdf_file = docs_dir / "Predictive_Smart_Irrigation_Analysis.pdf"
    
    if not md_file.exists():
        print(f"Markdown file not found: {md_file}")
        sys.exit(1)
    
    print("Converting markdown to Word and PDF formats...")
    
    # Convert to Word
    try:
        markdown_to_word(md_file, word_file)
    except ImportError:
        print("python-docx not installed. Install with: pip install python-docx")
    except Exception as e:
        print(f"Error creating Word document: {e}")
    
    # Convert to PDF via HTML
    try:
        markdown_to_pdf_html(md_file, html_file, pdf_file)
    except Exception as e:
        print(f"Error creating PDF: {e}")
    
    print("\nConversion complete!")
    print(f"Files created in: {docs_dir}")
    print(f"- Markdown: {md_file.name}")
    print(f"- Word: {word_file.name}")
    print(f"- HTML: {html_file.name}")
    print(f"- PDF: {pdf_file.name}")

if __name__ == "__main__":
    main()